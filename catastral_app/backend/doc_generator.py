from docx import Document
from docx.shared import Pt
from typing import Dict

def generar_docx(data: Dict, output_path: str = "certificado_generado.docx") -> str:
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    referencia = data.get("referencia_catastral", "[Referencia no disponible]")
    direccion = data.get("direccion", "[Dirección no disponible]")
    superficie = data.get("superficie_m2", 0)
    valor = data.get("valor_catastral", 0)
    titulares = data.get("titulares", [])
    doc.add_paragraph(f"1.- Urbana.- Solar sito en {direccion}, municipio de O Porriño (Pontevedra), con una superficie de {superficie:.2f} metros cuadrados.")
    doc.add_paragraph("Linda: [Linderos no disponibles en PDF]")
    doc.add_paragraph(f"\nValor: {valor:,.2f} €")
    doc.add_paragraph("\nVALOR DE REFERENCIA.- Yo la Notaria incorporo a la presente la certificación catastral de valor de referencia, que he obtenido por medios telemáticos, en el que no figura valor de referencia.")
    doc.add_paragraph(f"\nReferencia catastral: {referencia}")
    doc.add_paragraph("\nTitular/es:")
    for t in titulares:
        doc.add_paragraph(f"- {t['nombre']} ({t['nif']}): {t['porcentaje']}%", style='List Bullet')
    doc.save(output_path)
    return output_path
